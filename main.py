import base64
import json
import io
import os
import re
import zipfile
import olefile
import sqlite3
import time
from datetime import datetime
from PyPDF2 import PdfReader
from fastapi import FastAPI, UploadFile, File, Form
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse
from openai import OpenAI
import uvicorn
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from openpyxl import Workbook

app = FastAPI()

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True, 
    allow_methods=["*"],
    allow_headers=["*"],
)

OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")
client = OpenAI(api_key=OPENAI_API_KEY)

# Word 표 테두리 정밀 제어를 위한 함수
def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)
            for key, val in edge_data.items():
                element.set(qn('w:{}'.format(key)), str(val))

def init_db():
    conn = sqlite3.connect("hospital_ai.db")
    cursor = conn.cursor()
    cursor.execute("CREATE TABLE IF NOT EXISTS users (username TEXT PRIMARY KEY, password TEXT)")
    cursor.execute("INSERT OR IGNORE INTO users (username, password) VALUES ('admin', '1234')")
    cursor.execute("CREATE TABLE IF NOT EXISTS projects (project_id TEXT PRIMARY KEY, project_name TEXT, filename TEXT, plan_text TEXT, budget INTEGER)")
    cursor.execute("""
        CREATE TABLE IF NOT EXISTS minutes (
            id INTEGER PRIMARY KEY AUTOINCREMENT, project_id TEXT, meeting_type TEXT, store TEXT, 
            date TEXT, amount INTEGER, plan_task TEXT, time TEXT, location TEXT, attendees TEXT, 
            content TEXT, status TEXT, violation_reason TEXT, input_guide TEXT
        )
    """)
    cursor.execute("CREATE TABLE IF NOT EXISTS contracts (id INTEGER PRIMARY KEY AUTOINCREMENT, project_id TEXT, researcher_name TEXT, salary INTEGER, start_date TEXT, end_date TEXT)")
    conn.commit()
    conn.close()

init_db()

def extract_text(file_bytes, filename):
    text = ""
    ext = filename.lower().split('.')[-1]
    try:
        if ext == "pdf":
            reader = PdfReader(io.BytesIO(file_bytes))
            for page in reader.pages: text += page.extract_text() + "\n"
        elif ext == "hwp":
            f = io.BytesIO(file_bytes)
            if olefile.isOleFile(f):
                ole = olefile.OleFileIO(f)
                if ole.exists('PrvText'): text = ole.openstream('PrvText').read().decode('utf-16le', errors='ignore')
        elif ext == "hwpx":
            with zipfile.ZipFile(io.BytesIO(file_bytes)) as z:
                for item in z.namelist():
                    if item.endswith('.xml'): text += z.read(item).decode('utf-8', errors='ignore')
        elif ext == "docx":
            doc_obj = Document(io.BytesIO(file_bytes))
            for p in doc_obj.paragraphs: text += p.text + "\n"
            for table in doc_obj.tables:
                for row in table.rows:
                    for cell in row.cells: text += cell.text + " "
                    text += "\n"
    except Exception as e:
        print(f"추출 에러: {e}")
    return text

@app.post("/setup-demo-data")
async def setup_demo_data():
    project_id = "proj_seoul_st_mary_demo"
    project_name = "(예시)서울성모병원 연구과제"
    budget = 100000000
    conn = sqlite3.connect("hospital_ai.db")
    cursor = conn.cursor()
    
    cursor.execute("DELETE FROM projects WHERE project_id = ?", (project_id,))
    cursor.execute("INSERT INTO projects (project_id, project_name, filename, plan_text, budget) VALUES (?, ?, '서울성모병원_국책과제_사업계획서.hwp', '서울성모병원 연구과제 데이터셋', ?)", (project_id, project_name, budget))
    
    cursor.execute("DELETE FROM contracts WHERE project_id = ?", (project_id,))
    cursor.execute("INSERT INTO contracts (project_id, researcher_name, salary, start_date, end_date) VALUES (?, '김기훈', 2800000, '2022-06-01', '2027-12-31')", (project_id,))
    cursor.execute("INSERT INTO contracts (project_id, researcher_name, salary, start_date, end_date) VALUES (?, '홍준기', 3200000, '2022-01-01', '2022-12-31')", (project_id,))
    
    cursor.execute("DELETE FROM minutes WHERE project_id = ?", (project_id,))
    
    demo_content = """1. 임상시험 방법에 대한 논의
   가. 5년치의 데이터를 다 쓰는 방안과 1회성의 데이터로 재발률을 예측하는 방안에 대한 비교 논의
   나. 가장 최근 검사일 기준의 1회성 데이터로 임상시험 진행하기로 결정
2. 임상시험용 SW 수정에 관한 논의
   가. 초음파 팝업 전체 삭제
   나. 재발률 퍼센트로 표기하는 UI 수정

[제출용 자동 기안 공문]
수신: 서울성모병원 연구부
제목: 국책과제 연구활동비(회의비) 지출 결의 상신

1. 귀 부서의 노고에 감사드립니다.
2. 관련근거: 서울성모병원 연구비 관리 지침
3. 위 관련하여, 위 연구개발과제의 성공적 수행을 위하여 첨부와 같이 기 집행된 연구비 법인카드 지출 내역(회의비)을 결의하오니, 검토 후 승인하여 주시기 바랍니다.

AI센터장"""
    
    cursor.execute("""
        INSERT INTO minutes (project_id, meeting_type, store, date, amount, plan_task, time, location, attendees, content, status, violation_reason, input_guide)
        VALUES (?, 'conference', '청주경복궁', '2024-04-08', 283000, 
        '임상시험계획서 검토',
        '2024년 4월 8일 19:09 ~ 20:24', '서울성모병원 서관 9층 회의실', '내부: 김기훈, 이유진 외 4명 / 외부: 이강빈(아이도트)',
        ?, 'normal', '정상', '비목: 직접비 > 연구활동비 > 회의비\\n결제수단: 연구비 법인카드\\n금액: 283,000원')
    """, (project_id, demo_content))
    
    conn.commit()
    conn.close()
    return {"status": "success", "project_id": project_id, "project_name": project_name, "budget": budget}

@app.post("/upload-plan")
async def upload_plan(project_id: str = Form(...), project_name: str = Form(...), budget: int = Form(...), plan: UploadFile = File(...)):
    file_bytes = await plan.read()
    full_text = extract_text(file_bytes, plan.filename)
    plan_text_sliced = full_text[:15000] if full_text.strip() else "텍스트를 추출할 수 없거나 비어있는 파일입니다."
    conn = sqlite3.connect("hospital_ai.db")
    cursor = conn.cursor()
    cursor.execute("INSERT OR REPLACE INTO projects (project_id, project_name, filename, plan_text, budget) VALUES (?, ?, ?, ?, ?)", (project_id, project_name, plan.filename, plan_text_sliced, budget))
    conn.commit()
    conn.close()
    return {"status": "success"}

@app.post("/upload-receipt")
async def process_receipt(project_id: str = Form(...), receipts: list[UploadFile] = File(...), category: str = Form("conference")):
    conn = sqlite3.connect("hospital_ai.db")
    cursor = conn.cursor()
    cursor.execute("SELECT plan_text FROM projects WHERE project_id = ?", (project_id,))
    row = cursor.fetchone()
    if not row:
        conn.close()
        return {"error": "사업계획서가 먼저 등록되어야 합니다."}
    plan_text = row[0]
    output_results = []
    
    for receipt in receipts:
        receipt_contents = await receipt.read()
        base64_image = base64.b64encode(receipt_contents).decode('utf-8')
        
        common_instruction = "너는 서울성모병원 연구부의 수석 행정관이야. 제공된 [사업계획서]를 기반으로 과업 연계성을 입증해라."

        if category == "equipment":
            system_prompt = f"""{common_instruction}
            - 시설장비비 규정을 심사하고 대금 지급 공문을 작성해라.
            - 공문 작성 시 반드시 1.인사말, 2.관련근거, 3.요청사항(가,나,다 항목 포함) 형태의 표준 양식을 준수해라.
            {{
                "store": "공급사명", "date": "취득일자", "amount": 10000, "plan_task": "도입 타당성",
                "settlement_status": "normal", "violation_reason": "없음", "system_input_guide": "비목: 직접비 > 연구시설장비비",
                "official_memo": {{
                    "receiver": "서울성모병원 연구부",
                    "title": "국책과제 연구시설장비 대금 지급 요청",
                    "body": "1. 귀 부서의 노고에 감사드립니다.\\n2. 관련근거: 서울성모병원 연구비 지침\\n3. 위 관련하여, 다음과 같이 장비 대금 지급을 요청하오니 검토 후 승인하여 주시기 바랍니다.\\n 가. 안건: ...\\n 나. 금액: ...\\n 다. 사유: ...",
                    "sender": "AI센터장"
                }},
                "minutes": {{ "time": "검수일자", "location": "장소", "attendees": "검수원", "content": "검수조서 본문" }}
            }}"""
        elif category == "material":
            system_prompt = f"""{common_instruction}
            - 재료비 규정을 심사하고 대금 지급 공문을 작성해라.
            - 공문 작성 시 반드시 1.인사말, 2.관련근거, 3.요청사항(가,나,다 항목 포함) 형태의 표준 양식을 준수해라.
            {{
                "store": "공급사명", "date": "입고일자", "amount": 10000, "plan_task": "연계 타당성",
                "settlement_status": "normal", "violation_reason": "없음", "system_input_guide": "비목: 직접비 > 연구재료비",
                "official_memo": {{
                    "receiver": "서울성모병원 연구부",
                    "title": "국책과제 연구재료비 대금 지급 요청",
                    "body": "1. 귀 부서의 노고에 감사드립니다.\\n2. 관련근거: ...\\n3. 위 관련하여, 다음과 같이 재료비 지급을 요청하오니 승인하여 주시기 바랍니다.\\n 가. 안건: ...\\n 나. 금액: ...\\n 다. 사유: ...",
                    "sender": "AI센터장"
                }},
                "minutes": {{ "time": "검수일자", "location": "장소", "attendees": "검수원", "content": "검수 본문" }}
            }}"""
        else:
            system_prompt = f"""{common_instruction}
            - 회의비 규정을 심사해라. 회의비는 이미 '법인카드'로 결제된 건이므로 절대 '업체에 대금 지급을 요청한다'는 말을 쓰지 마라.
            - 반드시 JSON 내 `minutes.content`에 제공된 사업계획서를 바탕으로 회의록 양식에 들어갈 '1. 임상시험 논의 가. ... 2. SW 수정 논의 가. ...' 와 같이 아주 구체적이고 전문적인 [회의록 상세 본문]을 길게 창작해라.
            - 공문 작성 시 반드시 1.인사말, 2.관련근거, 3.요청사항(가,나,다 항목 포함) 형태의 표준 양식을 준수해라.
            {{
                "store": "가맹점명", "date": "결제일시", "amount": 10000, "plan_task": "회의 목적 (예: 임상시험계획서 검토 등)",
                "settlement_status": "normal / caution / invalid", "violation_reason": "판정 사유",
                "system_input_guide": "비목: 연구활동비 > 회의비",
                "official_memo": {{
                    "receiver": "서울성모병원 연구부",
                    "title": "국책과제 회의비 지출 결의(법인카드) 상신",
                    "body": "1. 귀 부서의 노고에 감사드립니다.\\n2. 관련근거: 서울성모병원 연구비 지침\\n3. 위 관련하여, 다음과 같이 기 집행된 카드 지출 내역을 결의하오니 승인하여 주시기 바랍니다.\\n 가. 회의목적: ...\\n 나. 집행금액: ...",
                    "sender": "AI센터장"
                }},
                "minutes": {{ 
                    "time": "202X년 X월 X일 19:00 ~ 20:30", 
                    "location": "서울성모병원 회의실", 
                    "attendees": "내부: OOO / 외부: OOO", 
                    "content": "1. 임상시험 방법에 대한 논의\\n 가. 5년치 데이터 사용 방안...\\n2. SW 수정 논의\\n 가. 초음파 팝업 삭제..." 
                }}
            }}"""

        try:
            response = client.chat.completions.create(
                model="gpt-4o",
                response_format={ "type": "json_object" },
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": [
                        {"type": "text", "text": f"[사업계획서 발췌]\n{plan_text}\n\n[영수증 이미지 분석]"},
                        {"type": "image_url", "image_url": {"url": f"data:image/jpeg;base64,{base64_image}"}}
                    ]}
                ]
            )
            result = json.loads(response.choices[0].message.content)
            m = result.get('minutes', {})
            memo = result.get('official_memo', {})
            
            raw_amount = str(result.get("amount", "0"))
            cleaned_amount = "".join(filter(str.isdigit, raw_amount))
            amount_val = int(cleaned_amount) if cleaned_amount else 0
            
            result['amount'] = amount_val
            result['store'] = result.get('store', '알 수 없음')
            result['plan_task'] = result.get('plan_task', '과업 연계성 검토 수행')
            result['settlement_status'] = result.get('settlement_status', 'normal')
            result['violation_reason'] = result.get('violation_reason', '정상')
            result['system_input_guide'] = result.get('system_input_guide', '비목: 연구활동비 > 회의비')
            
            memo_text = f"수신: {memo.get('receiver', '서울성모병원 연구부')}\n제목: {memo.get('title', '')}\n\n{memo.get('body', '')}\n\n{memo.get('sender', '')}"
            final_content = f"{m.get('content', '')}\n\n[제출용 자동 기안 공문]\n{memo_text}"
            
            cursor.execute("""
                INSERT INTO minutes (project_id, meeting_type, store, date, amount, plan_task, time, location, attendees, content, status, violation_reason, input_guide)
                VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
            """, (
                project_id, category, result.get('store'), result.get('date'), result.get('amount'),
                result.get('plan_task'), m.get('time'), m.get('location'), m.get('attendees'), final_content,
                result.get('settlement_status'), result.get('violation_reason'), result.get('system_input_guide')
            ))
            
            result['minute_id'] = cursor.lastrowid
            result['final_content'] = final_content
            output_results.append(result)
        except Exception as e:
            output_results.append({"error": str(e)})
            
    conn.commit()
    conn.close()
    return output_results

@app.post("/execute-expense-rpa")
async def execute_expense_rpa(minute_id: int = Form(...)):
    conn = sqlite3.connect("hospital_ai.db")
    cursor = conn.cursor()
    cursor.execute("SELECT store, amount, meeting_type FROM minutes WHERE id = ?", (minute_id,))
    row = cursor.fetchone()
    if not row:
        conn.close()
        return {"status": "error", "message": "결재 데이터를 찾을 수 없습니다."}
        
    store_name, amount, meeting_type = row
    cursor.execute("UPDATE minutes SET status = 'settled' WHERE id = ?", (minute_id,))
    conn.commit()
    conn.close()
    
    target_dept = "서울성모병원 연구부"
    if meeting_type == 'conference':
        message = f"🎉 [지출 결의 상신 완료]\n{target_dept}로 지출 결의서 및 협조 공문이 전송되었습니다."
    else:
        message = f"🎉 [대금 지급 상신 완료]\n{target_dept}로 지출 결의서 및 협조 공문이 전송되었습니다.\n승인 완료 시 '{store_name}'(으)로 대금({amount:,}원)이 자동 송금됩니다."
    return {"status": "success", "message": message}

@app.post("/sync-ezbaro")
async def sync_ezbaro(project_id: str = Form(...), billing_month: str = Form(...)):
    conn = sqlite3.connect("hospital_ai.db")
    cursor = conn.cursor()
    cursor.execute("SELECT researcher_name, end_date, salary FROM contracts WHERE project_id = ?", (project_id,))
    contracts = cursor.fetchall()
    
    if not contracts:
        conn.close()
        return {"status": "error", "message": "등록된 연구원 데이터가 없습니다."}
        
    expired_people = []
    active_payroll_amount = 0
    
    for name, end_date_str, salary in contracts:
        try:
            contract_end = datetime.strptime(end_date_str, "%Y-%m-%d")
            billing_date = datetime.strptime(f"{billing_month}-01", "%Y-%m-%d")
            if contract_end < billing_date:
                expired_people.append(f"{name}(만료일: {end_date_str})")
            else:
                active_payroll_amount += salary
        except: pass
            
    if expired_people:
        conn.close()
        return {"status": "intercepted", "message": "계약기간 경과 인원이 있어 자동 결재를 차단했습니다.", "details": expired_people}
        
    try:
        time.sleep(1.5) 
        cursor.execute("""
            INSERT INTO minutes (project_id, meeting_type, store, date, amount, plan_task, time, location, attendees, content, status, violation_reason, input_guide)
            VALUES (?, 'labor', '인사행정시스템', ?, ?, '연구인건비 자동 계상', ?, '원내 인사 전산망', '계약 연구원 전체', '자동 집행 완료', 'settled', '정상', 'RPA 상신 완료')
        """, (project_id, billing_month, active_payroll_amount, billing_month))
        conn.commit()
        conn.close()
        return {"status": "success", "message": f"🎉 [{billing_month}] 총 {active_payroll_amount:,}원의 인건비 지급 기안이 서울성모병원 연구부로 상신되었습니다."}
    except Exception as e:
        if conn: conn.close()
        return {"status": "error", "message": str(e)}

@app.get("/project-stats/{project_id}")
async def get_project_stats(project_id: str):
    conn = sqlite3.connect("hospital_ai.db")
    cursor = conn.cursor()
    cursor.execute("SELECT project_name, budget FROM projects WHERE project_id = ?", (project_id,))
    proj_row = cursor.fetchone()
    if not proj_row:
        conn.close()
        return {"budget": 0, "total_spent": 0, "remaining": 0, "normal": 0, "caution": 0, "invalid": 0}
    p_name, budget = proj_row
    cursor.execute("SELECT amount, status FROM minutes WHERE project_id = ?", (project_id,))
    minutes_rows = cursor.fetchall()
    conn.close()
    total_spent = sum(row[0] for row in minutes_rows if row[0])
    return {
        "project_name": p_name, "budget": budget, "total_spent": total_spent, "remaining": budget - total_spent,
        "normal": sum(1 for r in minutes_rows if r[1] in ['normal', 'settled']),
        "caution": sum(1 for r in minutes_rows if r[1] == 'caution'),
        "invalid": sum(1 for r in minutes_rows if r[1] == 'invalid'),
        "total_count": len(minutes_rows)
    }

@app.get("/export-excel")
async def export_excel(project_id: str):
    conn = sqlite3.connect("hospital_ai.db")
    cursor = conn.cursor()
    cursor.execute("SELECT meeting_type, store, date, amount, plan_task, status, violation_reason FROM minutes WHERE project_id = ? AND meeting_type != 'real' ORDER BY id ASC", (project_id,))
    rows = cursor.fetchall()
    conn.close()
    
    wb = Workbook()
    ws = wb.active
    ws.title = "ezbaro_Bulk_Template"
    headers = ["비목", "세부비목", "증빙일자(납품일)", "가맹점/공급처", "총지출액", "공급가액", "부가세", "적요(집행타당성 사유)", "감사결과"]
    ws.append(headers)
    
    for row in rows:
        m_type, store, date_str, amount, plan_task, status, violation = row
        amount = amount if amount else 0
        supply_value = int(amount / 1.1) if m_type != "labor" else amount
        vat = amount - supply_value if m_type != "labor" else 0
        ws.append([m_type, "세부비목", date_str, store, amount, supply_value, vat, plan_task, status])
        
    file_stream = io.BytesIO()
    wb.save(file_stream)
    file_stream.seek(0)
    return StreamingResponse(file_stream, media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", headers={"Content-Disposition": f"attachment; filename={project_id}_Upload.xlsx"})

@app.get("/export-word")
async def export_word(project_id: str):
    conn = sqlite3.connect("hospital_ai.db")
    cursor = conn.cursor()
    cursor.execute("SELECT project_name FROM projects WHERE project_id = ?", (project_id,))
    proj_row = cursor.fetchone()
    project_name = proj_row[0] if proj_row else "서울성모병원 원내과제"
    
    cursor.execute("SELECT meeting_type, store, amount, plan_task, time, location, attendees, content FROM minutes WHERE project_id = ? ORDER BY id ASC", (project_id,))
    rows = cursor.fetchall()
    conn.close()
    
    doc = Document()
    
    for idx, row in enumerate(rows):
        m_type, store, amount, plan_task, meeting_time, location, attendees, content = row
        
        if m_type == "conference":
            title = doc.add_paragraph()
            r = title.add_run("회   의   록")
            r.font.name = 'Malgun Gothic'
            r.font.size = Pt(18)
            r.bold = True
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            doc.add_paragraph("1. 과제 개요")
            t1 = doc.add_table(rows=5, cols=4)
            t1.style = 'Table Grid'
            t1.cell(0,0).text = "사 업 명"
            t1.cell(0,1).text = "국가연구개발사업(원내과제)"
            t1.cell(0,2).text = "주관기관"
            t1.cell(0,3).text = "서울성모병원"
            t1.cell(1,0).text = "연구개발과제명"
            t1.cell(1,1).merge(t1.cell(1,3)).text = project_name
            t1.cell(2,0).text = "연구책임자"
            t1.cell(2,1).text = "소속: 임상과"
            t1.cell(2,2).text = "성명:"
            t1.cell(2,3).text = "김기훈 (서명)"
            t1.cell(3,0).text = "전문기관명"
            t1.cell(3,1).text = "연구부"
            t1.cell(3,2).text = "과제번호"
            t1.cell(3,3).text = project_id
            t1.cell(4,0).text = "전체 연구기간"
            t1.cell(4,1).merge(t1.cell(4,3)).text = "2024. 1. 1. ~ 2027. 12. 31."
            
            doc.add_paragraph("\n2. 회의 내용")
            t2 = doc.add_table(rows=7, cols=2)
            t2.style = 'Table Grid'
            t2.columns[0].width = Inches(1.5)
            t2.columns[1].width = Inches(5.0)
            
            t2.cell(0,0).text = "회 의 일 시"
            t2.cell(0,1).text = meeting_time if meeting_time else ""
            t2.cell(1,0).text = "회 의 장 소"
            t2.cell(1,1).text = location if location else ""
            t2.cell(2,0).text = "참 석 자"
            t2.cell(2,1).text = attendees if attendees else ""
            t2.cell(3,0).text = "카드 사용처"
            t2.cell(3,1).text = store if store else ""
            t2.cell(4,0).text = "집 행 금 액"
            t2.cell(4,1).text = f"금{amount:,}원" if amount else ""
            t2.cell(5,0).text = "회 의 목 적"
            t2.cell(5,1).text = plan_task if plan_task else ""
            t2.cell(6,0).text = "회 의 내 용"
            
            clean_content = content.split("[제출용 자동 기안 공문]")[0].strip() if content else ""
            t2.cell(6,1).text = clean_content
            
            if idx < len(rows) - 1:
                doc.add_page_break()
        else:
            doc.add_paragraph(f"[{m_type.upper()}] 정산 증빙서 - {store}")
            doc.add_paragraph(f"금액: {amount:,}원")
            clean_content = content.split("[제출용 자동 기안 공문]")[0].strip() if content else ""
            doc.add_paragraph(f"내용:\n{clean_content}")
            if idx < len(rows) - 1:
                doc.add_page_break()
                
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return StreamingResponse(file_stream, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document", headers={"Content-Disposition": f"attachment; filename={project_id}_Meeting_Minutes.docx"})

# 🛠️ 대대적 수정: 공문 양식을 HWP 내부 양식과 똑같이 정밀 모델링하여 파일 생성
@app.get("/export-memo/{minute_id}")
async def export_memo(minute_id: int):
    conn = sqlite3.connect("hospital_ai.db")
    cursor = conn.cursor()
    cursor.execute("SELECT content, store, amount, plan_task, date FROM minutes WHERE id = ?", (minute_id,))
    row = cursor.fetchone()
    conn.close()

    if not row:
        return {"error": "데이터를 찾을 수 없습니다."}

    full_content, store, amount, plan_task, doc_date = row
    memo_parts = full_content.split("[제출용 자동 기안 공문]")
    memo_lines = memo_parts[1].strip().split('\n') if len(memo_parts) > 1 else []

    # 파싱 안전장치
    title_text = "국책과제 지출 결의 상신"
    for line in memo_lines:
        if line.startswith("제목:"):
            title_text = line.replace("제목:", "").strip()

    doc = Document()
    
    # 여백 설정 (공문 표준 여백)
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # 1. 상단 발신 기관명 헤더
    p_header = doc.add_paragraph()
    p_header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_header = p_header.add_run("서 울 성 모 병 원  연 구 부")
    run_header.font.name = 'Malgun Gothic'
    run_header.font.size = Pt(22)
    run_header.font.bold = True
    run_header.font.color.rgb = RGBColor(15, 23, 42)
    p_header.paragraph_format.space_after = Pt(20)

    # 2. 문서 정보 격자 메타데이터 테이블 (테두리 없음 우아한 배치)
    meta_table = doc.add_table(rows=3, cols=2)
    meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta_table.autofit = False
    
    meta_rows = [
        (f"문서번호 : AI센터-{datetime.now().year}-정산", "선 결 재 : [ 생 략 ]"),
        ("수    신 : 서울성모병원 연구부", f"시행일자 : {doc_date if doc_date and '-' in doc_date else datetime.now().strftime('%Y.%m.%d')}"),
        (f"제    목 : {title_text}", "")
    ]
    
    for idx, (left, right) in enumerate(meta_rows):
        row = meta_table.rows[idx]
        cell_l, cell_r = row.cells[0], row.cells[1]
        cell_l.width = Inches(4.5)
        cell_r.width = Inches(2.5)
        
        pl = cell_l.paragraphs[0]
        pl.paragraph_format.space_after = Pt(4)
        rl = pl.add_run(left)
        rl.font.name = 'Malgun Gothic'
        rl.font.size = Pt(11)
        if "제    목" in left:
            rl.font.bold = True
            
        pr = cell_r.paragraphs[0]
        pr.paragraph_format.space_after = Pt(4)
        pr.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        rr = pr.add_run(right)
        rr.font.name = 'Malgun Gothic'
        rr.font.size = Pt(11)

    # 구분선 추가
    p_sep = doc.add_paragraph()
    p_sep.paragraph_format.space_before = Pt(6)
    p_sep.paragraph_format.space_after = Pt(18)
    run_sep = p_sep.add_run("━" * 55)
    run_sep.font.color.rgb = RGBColor(200, 200, 200)

    # 3. 공문 표준 본문
    body_texts = [
        "1. 귀 부서의 무궁한 발전을 기원합니다.",
        "2. 관련근거: 서울성모병원 연구비 집행 지침 및 기준 규정",
        "3. 위 관련하여, 성공적인 과업 완수 및 목표 연구 성과 달성을 위하여 기 집행된 내역을 다음과 같이 보고 및 결의하오니 검토 후 승인하여 주시기 바랍니다."
    ]
    
    for text in body_texts:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = 1.3
        run = p.add_run(text)
        run.font.name = 'Malgun Gothic'
        run.font.size = Pt(11)

    # 4. 세부 정산 명세 테이블 (HWP 형태의 가, 나, 다 구조화 블록)
    details = [
        ("가. 정산비목 :", "국่าย연구개발과제 연구비 (지출 증빙)"),
        ("나. 가 맹 점 :", f"{store}"),
        ("다. 결제금액 :", f"금 {amount:,}원 (정액 수령 및 실집행 규격)"),
        ("라. 집행목적 :", f"{plan_task}")
    ]
    
    detail_table = doc.add_table(rows=4, cols=2)
    detail_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for idx, (label, val) in enumerate(details):
        row = detail_table.rows[idx]
        c1, c2 = row.cells[0], row.cells[1]
        c1.width = Inches(1.2)
        c2.width = Inches(5.3)
        
        p1 = c1.paragraphs[0]
        p1.paragraph_format.space_after = Pt(4)
        r1 = p1.add_run(label)
        r1.font.name = 'Malgun Gothic'
        r1.font.size = Pt(11)
        r1.font.bold = True
        
        p2 = c2.paragraphs[0]
        p2.paragraph_format.space_after = Pt(4)
        r2 = p2.add_run(val)
        r2.font.name = 'Malgun Gothic'
        r2.font.size = Pt(11)

    doc.add_paragraph("\n4. 첨부서류: 전자 영수증 및 증빙 파일 조서 1부.  끝.")
    doc.paragraphs[-1].runs[0].font.name = 'Malgun Gothic'
    doc.paragraphs[-1].runs[0].font.size = Pt(11)

    # 5. 하단 발신인 직인 란
    p_sender = doc.add_paragraph()
    p_sender.paragraph_format.space_before = Pt(45)
    p_sender.paragraph_format.space_after = Pt(40)
    p_sender.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sender = p_sender.add_run("서 울 성 모 병 원  A I Center 장  [ 직 인 생 략 ]")
    run_sender.font.name = 'Malgun Gothic'
    run_sender.font.size = Pt(15)
    run_sender.font.bold = True
    run_sender.font.color.rgb = RGBColor(30, 41, 59)

    # 6. 하단 테두리 정보바 (기관 메타 하단정보바 재현)
    footer_table = doc.add_table(rows=1, cols=1)
    footer_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    f_cell = footer_table.rows[0].cells[0]
    f_cell.width = Inches(7.0)
    set_cell_border(f_cell, top={"sz": 6, "val": "single", "color": "A0A0A0"})

    fp = f_cell.paragraphs[0]
    fp.paragraph_format.space_before = Pt(6)
    frun = fp.add_run("우 06591 서울특별시 서초구 반포대로 222  /  홈페이지: https://www.cmcseoul.or.kr  /  전화: 02-2258-6114  /  공개")
    frun.font.name = 'Malgun Gothic'
    frun.font.size = Pt(8.5)
    frun.font.color.rgb = RGBColor(120, 120, 120)

    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return StreamingResponse(file_stream, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document", headers={"Content-Disposition": f"attachment; filename=Official_Memo_{minute_id}.docx"})

if __name__ == "__main__":
    uvicorn.run("main:app", host="127.0.0.1", port=8000, reload=True)
